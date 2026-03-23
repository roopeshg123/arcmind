"""
Generates ArcMind_Project_Overview.docx
Focused on: Architecture · How to Run · UI Buttons Guide
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── Colour palette ──────────────────────────────────────────────────────────
BRAND_BLUE   = RGBColor(0x3B, 0x4F, 0xD8)
GREEN_ACCENT = RGBColor(0x4C, 0xAF, 0x88)
MID_GREY     = RGBColor(0x88, 0x88, 0x99)
WHITE        = RGBColor(0xFF, 0xFF, 0xFF)
TEXT_DARK    = RGBColor(0x1A, 0x1A, 0x1A)
TEXT_BODY    = RGBColor(0x33, 0x33, 0x44)
DARK_NAVY    = RGBColor(0x1A, 0x1A, 0x2E)

# ─── Helpers ─────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, color="CCCCCC"):
    """Add thin borders to a table cell."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")
        el.set(qn("w:color"), color)
        borders.append(el)
    tcPr.append(borders)


def heading(doc, text, level, color=BRAND_BLUE):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.color.rgb = color
    return h


def para(doc, text, bold=False, color=None, size=11,
         before=0, after=6, italic=False, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)
    if align:
        p.alignment = align
    r = p.add_run(text)
    r.bold   = bold
    r.italic = italic
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    return p


def bullet(doc, text, prefix=None, prefix_color=BRAND_BLUE, after=4):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(after)
    if prefix:
        r1 = p.add_run(prefix)
        r1.bold = True
        r1.font.color.rgb = prefix_color
        r1.font.size = Pt(10.5)
        r2 = p.add_run(text)
        r2.font.size = Pt(10.5)
    else:
        r = p.add_run(text)
        r.font.size = Pt(10.5)
    return p


def callout(doc, lines, bg="E8ECFD", fg=TEXT_DARK, mono=False):
    """Single-cell shaded box."""
    tbl  = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, bg)
    cell.width = Inches(6.5)
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        p.paragraph_format.left_indent  = Pt(10)
        p.paragraph_format.right_indent = Pt(10)
        r = p.add_run(line)
        r.font.size      = Pt(10)
        r.font.color.rgb = fg
        if mono:
            r.font.name = "Courier New"
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def code_block(doc, lines):
    """Dark code/command block."""
    callout(doc, lines, bg="1E1E2E", fg=RGBColor(0xCC, 0xCC, 0xEE), mono=True)


def divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run("─" * 90)
    r.font.size      = Pt(7)
    r.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)


# ─── Build document ──────────────────────────────────────────────────────────

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)


# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

cp = doc.add_paragraph()
cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
cr = cp.add_run("ArcMind")
cr.font.size = Pt(54)
cr.font.bold = True
cr.font.color.rgb = BRAND_BLUE

cp2 = doc.add_paragraph()
cp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
cr2 = cp2.add_run("Enterprise AI Assistant for Arc")
cr2.font.size = Pt(20)
cr2.font.color.rgb = TEXT_BODY

doc.add_paragraph()

cp3 = doc.add_paragraph()
cp3.alignment = WD_ALIGN_PARAGRAPH.CENTER
cr3 = cp3.add_run("Chat with your Arc Documentation  ·  Jira Tickets  ·  Confluence Wiki")
cr3.font.size   = Pt(12)
cr3.font.italic = True
cr3.font.color.rgb = MID_GREY

doc.add_paragraph()
doc.add_paragraph()

# dark tagline box
tbl_cover = doc.add_table(rows=1, cols=1)
tbl_cover.alignment = WD_TABLE_ALIGNMENT.CENTER
cc = tbl_cover.cell(0, 0)
set_cell_bg(cc, "1A1A2E")
cc.width = Inches(5.5)
tpc = cc.paragraphs[0]
tpc.alignment = WD_ALIGN_PARAGRAPH.CENTER
tpc.paragraph_format.space_before = Pt(18)
tpc.paragraph_format.space_after  = Pt(18)
tpc.paragraph_format.left_indent  = Pt(14)
tpc.paragraph_format.right_indent = Pt(14)
tcr = tpc.add_run(
    '"Ask questions in plain English.\n'
    'Get accurate, sourced answers instantly\n'
    '— powered by OpenAI GPT-4.1 + RAG."'
)
tcr.font.color.rgb = RGBColor(0xEC, 0xEC, 0xEC)
tcr.font.size   = Pt(12)
tcr.font.italic = True

doc.add_paragraph()

cbot = doc.add_paragraph()
cbot.alignment = WD_ALIGN_PARAGRAPH.CENTER
cbr = cbot.add_run("Version 2.1.0   ·   March 2026   ·   github.com/roopeshg123/arcmind")
cbr.font.size      = Pt(9)
cbr.font.color.rgb = MID_GREY

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — PROJECT OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════

heading(doc, "1.  Project Overview", 1)

para(doc,
    "ArcMind is an AI-powered chat assistant that reads your Arc HTML documentation, "
    "Jira ticket history, and Confluence wiki pages, then answers questions instantly "
    "through a browser-based chat interface. It is a private GPT that only knows "
    "about YOUR application and YOUR team's knowledge.",
    size=11, after=8)

para(doc,
    "It uses RAG (Retrieval-Augmented Generation) — a technique that retrieves the "
    "most relevant document chunks at query time and feeds them to GPT-4.1 as context, "
    "so every answer is grounded in your actual content, not generic AI knowledge.",
    size=11, after=8)

callout(doc, [
    "Example",
    "",
    '  You ask   :  "How do I configure SFTP authentication in Arc?"',
    "",
    '  ArcMind   :  "Navigate to the SFTP Connector settings and set AuthScheme to',
    '               PublicKey. Set SSHClientCert to the path of your .pem key file..."',
    "",
    "  Sources   :  SFTP.html  ·  SFTP-Advanced-Properties.html  ·  ARC-1234",
], bg="E8ECFD")

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════════

heading(doc, "2.  Architecture", 1)

para(doc,
    "ArcMind has two phases — Indexing (builds the knowledge base once) and "
    "Query (runs every time a question is asked).",
    size=11, after=10)

# ── 2.1 Indexing ──────────────────────────────────────────────────────────────
heading(doc, "2.1  Indexing Phase  (Building the Knowledge Base)", 2)

code_block(doc, [
    "  Arc HTML Docs           Jira REST API v3       Confluence REST API v1",
    "       |                        |                          |",
    "       v                        v                          v",
    "  [BeautifulSoup]         [jira_client.py]      [confluence_client.py]",
    "  Extract clean text      Fetch all tickets      Fetch all wiki pages",
    "       |                        |                          |",
    "       +------------------------+--------------------------+",
    "                                |",
    "                         [chunking.py]",
    "               Docs: 1500 tokens/300 overlap",
    "               Jira:  600 tokens/100 overlap",
    "               Confluence: 800 tokens/150 overlap",
    "                                |",
    "               [OpenAI text-embedding-3-large]",
    "               Convert each chunk → 3072-number vector",
    "                                |",
    "                     [ChromaDB  +  BM25 index]",
    "               arcmind_docs · arcmind_jira · arcmind_confluence",
    "               chroma_db/  (local SQLite  +  bm25_*.pkl files)",
])

para(doc,
    "This runs once. After indexing you have a local vector database containing "
    "every chunk of your documentation, Jira tickets, and Confluence pages "
    "as searchable vectors — no internet required at query time.",
    size=10.5, after=10, color=RGBColor(0x44, 0x44, 0x55))

# ── 2.2 Query Pipeline ────────────────────────────────────────────────────────
heading(doc, "2.2  Query Pipeline  (Answering Every Question)", 2)

code_block(doc, [
    "  User types a question in the browser",
    "            |",
    "            v",
    "  [1] Connector Detection   — scan for Arc component names (SFTP, REST...)",
    "  [2] Query Router          — decide: docs-only / jira-only / hybrid",
    "  [3] Query Expansion       — GPT-4.1 generates 5 richer query variants",
    "  [4] Hybrid Retrieval      — ChromaDB vector search  +  BM25 keyword search",
    "                               across arcmind_docs, arcmind_jira, arcmind_confluence",
    "  [5] Cross-Encoder Rerank  — ms-marco-MiniLM-L-6-v2 (local, ~22 MB)",
    "                               re-scores every (question, chunk) pair jointly",
    "  [6] Prompt Builder        — assemble: Doc context + Confluence context + Jira context",
    "  [7] GPT-4.1               — generate answer, stream token-by-token via SSE",
    "  [8] Return + Log          — show answer with clickable sources, append to query_log.jsonl",
])

# ── 2.3 System Architecture Diagram ───────────────────────────────────────────
heading(doc, "2.3  System Diagram", 2)

code_block(doc, [
    "  ╔══════════════════════════════════════════════════════════════════╗",
    "  ║            BROWSER  —  static/index.html                        ║",
    "  ║  [ Index Docs ] [ Update Docs ] [ Index Jira ] [ Update Jira ]  ║",
    "  ║  [ Index Confluence ] [ Update Confluence ] [ Clear Chat ]      ║",
    "  ║  Status Badge  ·  Progress Bar  ·  Chat Window  ·  Input Box    ║",
    "  ╚════════════════════════╤═════════════════════════════════════════╝",
    "                           │  HTTP  (REST + SSE)",
    "  ╔════════════════════════▼═════════════════════════════════════════╗",
    "  ║           FASTAPI BACKEND  —  main.py  (port 8000)              ║",
    "  ║  POST /api/chat/stream   POST /api/ingest   GET /api/status     ║",
    "  ║  X-API-Key auth  ·  CORS  ·  asyncio.Lock (safe ingestion)      ║",
    "  ╚════╤═══════════════════╤════════════════════╤════════════════════╝",
    "       │                   │                    │",
    "  rag_engine.py      ingest_*.py          connectors/",
    "  Query pipeline     Crawl/Fetch          jira_client.py",
    "  7-step pipeline    Chunk & Embed        confluence_client.py",
    "       │                   │",
    "  ╔════▼═══════════════════▼════════════════════════════════════════╗",
    "  ║  ChromaDB  (chroma_db/)                                         ║",
    "  ║  arcmind_docs  ·  arcmind_jira  ·  arcmind_confluence           ║",
    "  ║  + bm25_docs.pkl  ·  bm25_jira.pkl  ·  bm25_confluence.pkl     ║",
    "  ╚═════════════════════════════════════════════════════════════════╝",
    "                           │",
    "              OpenAI API (cloud)",
    "              text-embedding-3-large  (indexing)",
    "              gpt-4.1 chat completions  (answering)",
])

# ── 2.4 Key Files ─────────────────────────────────────────────────────────────
heading(doc, "2.4  Key Files & What They Do", 2)

file_rows = [
    ("File / Folder",            "What it does"),
    ("main.py",                  "FastAPI server — all REST endpoints, auth middleware, CORS, async lock"),
    ("rag_engine.py",            "Runs the 7-step query pipeline, streaming output, query logging"),
    ("rag/retriever.py",         "Hybrid BM25 + vector search across all three collections"),
    ("rag/query_expander.py",    "Generates 5 query variants via GPT-4.1 for broader recall"),
    ("rag/reranker.py",          "Cross-encoder reranker — runs 100% locally, no API cost"),
    ("rag/prompt_builder.py",    "Assembles the GPT prompt with Docs + Confluence + Jira context"),
    ("rag/conversation_memory.py","Per-session server-side chat history (keyed by session_id)"),
    ("ingest/ingest_docs.py",    "Crawl Arc HTML docs or read local folder, chunk, embed, store"),
    ("ingest/ingest_jira.py",    "Full + incremental Jira ingestion with SHA-256 smart diff"),
    ("ingest/ingest_confluence.py","Full + CQL incremental Confluence ingestion with smart diff"),
    ("ingest/chunking.py",       "Token-aware text splitter — different sizes per source type"),
    ("vector_db/chroma_store.py","Manages ChromaDB collections + BM25 pickles + diff helpers"),
    ("connectors/jira_client.py","Jira REST API v3 with cursor pagination — fetches 100% of tickets"),
    ("connectors/confluence_client.py","Confluence REST API v1 — parses storage HTML, fixes wiki URLs"),
    ("static/index.html",        "Complete chat UI in vanilla HTML/CSS/JS — no framework needed"),
    ("chroma_db/",               "Auto-created local vector store (SQLite + BM25 pickle files)"),
    ("logs/query_log.jsonl",     "Every question, expanded queries, source counts, answer preview"),
    ("Dockerfile",               "Two-stage container build — reranker pre-downloaded at build time"),
    ("docker-compose.yml",       "Run the full app with: docker compose up --build"),
]

tbl_files = doc.add_table(rows=len(file_rows), cols=2)
tbl_files.style = "Table Grid"
tbl_files.alignment = WD_TABLE_ALIGNMENT.LEFT
for r_i, (col1, col2) in enumerate(file_rows):
    row = tbl_files.rows[r_i]
    row.cells[0].width = Inches(2.3)
    row.cells[1].width = Inches(4.2)
    for c_i, txt in enumerate([col1, col2]):
        cell = row.cells[c_i]
        if r_i == 0:
            set_cell_bg(cell, "3B4FD8")
        elif r_i % 2 == 0:
            set_cell_bg(cell, "F0F2FD")
        else:
            set_cell_bg(cell, "FAFAFA")
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        p.paragraph_format.left_indent  = Pt(5)
        run = p.add_run(txt)
        run.font.size = Pt(9.5)
        if r_i == 0:
            run.bold = True
            run.font.color.rgb = WHITE
        elif c_i == 0:
            run.bold = True
            run.font.color.rgb = BRAND_BLUE
            run.font.name = "Courier New"

doc.add_paragraph()
doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — TECH STACK
# ══════════════════════════════════════════════════════════════════════════════

heading(doc, "3.  Technology Stack", 1)

tech_rows = [
    ("Technology",              "Purpose"),
    ("Python 3.12",             "Backend language — all AI/ML libraries are Python-first"),
    ("FastAPI + Uvicorn",       "Async REST API framework + ASGI server"),
    ("LangChain",               "RAG chains, prompt management, LLM integration, chat history"),
    ("OpenAI GPT-4.1",          "Generates answers from retrieved context (Chat Completions API)"),
    ("text-embedding-3-large",  "Converts text chunks to 3,072-dimension semantic vectors"),
    ("ChromaDB",                "Local SQLite-backed vector database — three collections"),
    ("rank-bm25",               "BM25 keyword search index — pairs with vectors for hybrid search"),
    ("sentence-transformers",   "Cross-encoder reranker (ms-marco-MiniLM-L-6-v2) — local, ~22 MB"),
    ("BeautifulSoup4 / lxml",   "Parses Arc HTML docs — strips tags, extracts clean text"),
    ("httpx",                   "Async HTTP client for Jira and Confluence REST API calls"),
    ("tiktoken",                "Token counter — keeps chunks within GPT context limits"),
    ("Jira REST API v3",        "Fetches all tickets with cursor pagination"),
    ("Confluence REST API v1",  "Fetches wiki pages by space key; CQL for incremental sync"),
    ("Docker / Podman",         "One-command containerised deployment: docker compose up --build"),
    ("python-dotenv",           "Loads .env secrets — keeps API keys out of source code"),
]

tbl_tech = doc.add_table(rows=len(tech_rows), cols=2)
tbl_tech.style = "Table Grid"
tbl_tech.alignment = WD_TABLE_ALIGNMENT.LEFT
for r_i, (col1, col2) in enumerate(tech_rows):
    row = tbl_tech.rows[r_i]
    row.cells[0].width = Inches(2.5)
    row.cells[1].width = Inches(4.0)
    for c_i, txt in enumerate([col1, col2]):
        cell = row.cells[c_i]
        if r_i == 0:
            set_cell_bg(cell, "3B4FD8")
        elif r_i % 2 == 0:
            set_cell_bg(cell, "F0F2FD")
        else:
            set_cell_bg(cell, "FAFAFA")
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        p.paragraph_format.left_indent  = Pt(5)
        run = p.add_run(txt)
        run.font.size = Pt(9.5)
        if r_i == 0:
            run.bold = True
            run.font.color.rgb = WHITE
        elif c_i == 0:
            run.bold = True
            run.font.color.rgb = BRAND_BLUE

doc.add_paragraph()
doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — HOW TO RUN THE APP
# ══════════════════════════════════════════════════════════════════════════════

heading(doc, "4.  How to Run the App", 1)

para(doc,
    "There are two ways to run ArcMind. Docker is recommended for team deployments "
    "— no Python setup needed on each machine. Manual setup is for local development.",
    size=11, after=10)

# ── Option A — Docker ─────────────────────────────────────────────────────────
heading(doc, "Option A  —  Docker  (Recommended)", 2)

para(doc, "Prerequisites: Docker Desktop installed and running.", size=10.5, bold=True, after=6)

setup_docker = [
    ("Step 1", "Clone the repository",
     ["git clone https://github.com/roopeshg123/arcmind.git", "cd arcmind"]),
    ("Step 2", "Create your .env file  —  open it and fill in your keys",
     ["copy .env.example .env",
      "",
      "# Minimum required inside .env:",
      "OPENAI_API_KEY=sk-...",
      "JIRA_URL=https://your-company.atlassian.net",
      "JIRA_EMAIL=you@company.com",
      "JIRA_API_TOKEN=your-atlassian-token",
      "CONFLUENCE_URL=https://your-company.atlassian.net",
      "CONFLUENCE_EMAIL=you@company.com",
      "CONFLUENCE_API_TOKEN=your-atlassian-token",
      "CONFLUENCE_SPACES=ENG,PROD"]),
    ("Step 3", "Build and start the application",
     ["docker compose up --build",
      "",
      "# First run downloads ~200 MB (reranker model pre-cached in image)",
      "# Subsequent starts are instant"]),
    ("Step 4", "Open the chat UI in your browser",
     ["http://localhost:8000"]),
    ("Step 5", "Index your knowledge base using the toolbar buttons",
     ["Click  [ Index Docs ]        — indexes your Arc HTML documentation",
      "Click  [ Index Jira ]        — indexes all Jira tickets",
      "Click  [ Index Confluence ]  — indexes all Confluence pages",
      "",
      "Watch the progress bar fill to 100%",
      "Status badge turns green:  Ready — 75,104 vectors  when complete"]),
    ("Step 6", "Start chatting",
     ["Type any question in the input box and press Enter"]),
]

for step_num, step_title, step_cmds in setup_docker:
    # Step label row
    tbl_step = doc.add_table(rows=1, cols=2)
    tbl_step.alignment = WD_TABLE_ALIGNMENT.LEFT
    left_cell = tbl_step.cell(0, 0)
    right_cell = tbl_step.cell(0, 1)
    left_cell.width  = Inches(0.9)
    right_cell.width = Inches(5.6)

    set_cell_bg(left_cell, "3B4FD8")
    lp = left_cell.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lp.paragraph_format.space_before = Pt(6)
    lp.paragraph_format.space_after  = Pt(6)
    lr = lp.add_run(step_num)
    lr.font.size = Pt(10)
    lr.font.bold = True
    lr.font.color.rgb = WHITE

    rp = right_cell.paragraphs[0]
    rp.paragraph_format.space_before = Pt(6)
    rp.paragraph_format.space_after  = Pt(6)
    rp.paragraph_format.left_indent  = Pt(8)
    rr = rp.add_run(step_title)
    rr.font.size = Pt(11)
    rr.font.bold = True
    rr.font.color.rgb = BRAND_BLUE

    # Command block beneath
    code_block(doc, ["  " + c for c in step_cmds])

doc.add_paragraph()

# ── Option B — Manual Python ──────────────────────────────────────────────────
heading(doc, "Option B  —  Manual Python Setup", 2)

para(doc, "Prerequisites: Python 3.12 and Git installed.", size=10.5, bold=True, after=6)

code_block(doc, [
    "# 1. Clone",
    "git clone https://github.com/roopeshg123/arcmind.git",
    "cd arcmind",
    "",
    "# 2. Virtual environment",
    "python -m venv .venv",
    r".venv\Scripts\Activate.ps1",
    "",
    "# 3. Install dependencies (CPU-only PyTorch — ~200 MB total)",
    "pip install -r requirements.txt",
    "",
    "# 4. Fill in .env  (same keys as Docker Option A above)",
    "notepad .env",
    "",
    "# 5. Start the server",
    "uvicorn main:app --reload --host 0.0.0.0 --port 8000",
    "",
    "# 6. Open browser",
    "http://localhost:8000",
])

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — USER INTERFACE & BUTTONS
# ══════════════════════════════════════════════════════════════════════════════

heading(doc, "5.  User Interface", 1)

para(doc,
    "ArcMind runs entirely in the browser. No installation or plugin is needed "
    "for end users — just open http://localhost:8000.",
    size=11, after=10)

# ── 5.1 UI Layout ─────────────────────────────────────────────────────────────
heading(doc, "5.1  Interface Layout", 2)

# Mock-up of the real UI (matches the screenshot)
code_block(doc, [
    "  ┌─────────────────────────────────────────────────────────────────────────────────┐",
    "  │  ArcMind                    ✓ Ready — 75,104 vectors (docs:2809 jira:71385 ...)  │  ← HEADER",
    "  ├─────────────────────────────────────────────────────────────────────────────────┤",
    "  │ [⚡Index Docs][Update Docs][🔗Index Jira][Update Jira][📄Index Conf][Update Conf][Clear Chat] │  ← TOOLBAR",
    "  ├─────────────────────────────────────────────────────────────────────────────────┤",
    "  │                                                                                 │",
    "  │   You:       How do I configure SFTP authentication?                            │",
    "  │                                                                                 │",
    "  │   ArcMind:   To configure SFTP authentication in Arc:                           │  ← CHAT",
    "  │              • Set AuthScheme to PublicKey                                      │  WINDOW",
    "  │              • Set SSHClientCert to your .pem file path                         │",
    "  │              • Set SSHClientCertType to PEMKEY_FILE                             │",
    "  │              Sources: SFTP.html · ARC-1234                                      │",
    "  │                                                                                 │",
    "  ├─────────────────────────────────────────────────────────────────────────────────┤",
    "  │  [ Ask anything about your Arc application...                        ]  [Send]  │  ← INPUT",
    "  │  Press Enter to send  ·  Shift+Enter for new line                               │",
    "  └─────────────────────────────────────────────────────────────────────────────────┘",
])

# ── 5.2 Status Badge ──────────────────────────────────────────────────────────
heading(doc, "5.2  Status Badge  (Top Right)", 2)

status_rows = [
    ("Badge State",              "What it means"),
    ("✓ Ready — 75,104 vectors\n(docs:2809 jira:71385 conf:910)",
     "Knowledge base is indexed and ready. Shows exact chunk counts per source."),
    ("Not ready — no docs indexed",
     "No content has been indexed yet. Click Index Docs / Jira / Confluence to start."),
    ("Checking...",
     "App just started and is checking ChromaDB collections."),
]

tbl_status = doc.add_table(rows=len(status_rows), cols=2)
tbl_status.style = "Table Grid"
tbl_status.alignment = WD_TABLE_ALIGNMENT.LEFT
for r_i, (col1, col2) in enumerate(status_rows):
    row = tbl_status.rows[r_i]
    row.cells[0].width = Inches(2.8)
    row.cells[1].width = Inches(3.7)
    for c_i, txt in enumerate([col1, col2]):
        cell = row.cells[c_i]
        if r_i == 0:
            set_cell_bg(cell, "3B4FD8")
        else:
            set_cell_bg(cell, "F9FFF9" if c_i == 0 else "FAFAFA")
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        p.paragraph_format.left_indent  = Pt(6)
        run = p.add_run(txt)
        run.font.size = Pt(9.5)
        if r_i == 0:
            run.bold = True
            run.font.color.rgb = WHITE
        elif c_i == 0:
            run.bold = True
            run.font.color.rgb = GREEN_ACCENT
            run.font.name = "Courier New"

doc.add_paragraph()

# ── 5.3 Toolbar Buttons ────────────────────────────────────────────────────────
heading(doc, "5.3  Toolbar Buttons — What Each Button Does", 2)

btn_rows = [
    ("Button",            "Colour",  "What it does",                                               "When to use"),
    ("Index Docs",        "Blue",    "Full re-index of all Arc HTML documentation from scratch. "
                                     "Crawls the docs URL or reads the local DOCS_DIR folder. "
                                     "Strips HTML, chunks, embeds, stores in arcmind_docs collection.",
                                     "First-time setup, or when you want a complete fresh index."),
    ("Update Docs",       "Grey",    "Smart incremental update using SHA-256 content hashing. "
                                     "Only new or changed pages are re-embedded. "
                                     "Unchanged pages are skipped entirely (zero API cost).",
                                     "After documentation has been updated — fast and cost-efficient."),
    ("Index Jira",        "Blue",    "Full re-index of all Jira issues in your project. "
                                     "Fetches tickets with cursor pagination via Jira REST API v3. "
                                     "Indexes each ticket body AND all comments separately.",
                                     "First-time setup or full Jira re-sync."),
    ("Update Jira",       "Grey",    "SHA-256 smart diff update for Jira. "
                                     "Only tickets that are new, edited, or have new comments are re-indexed. "
                                     "Also supports incremental sync for the last N hours.",
                                     "Daily/weekly keeps Jira knowledge base current cheaply."),
    ("Index Confluence",  "Blue",    "Full re-index of all Confluence pages in configured spaces. "
                                     "Fetches via Confluence REST API v1, parses storage-format HTML, "
                                     "corrects wiki URLs to be clickable in answers.",
                                     "First-time setup or after major Confluence restructuring."),
    ("Update Confluence", "Grey",    "Smart incremental update for Confluence using CQL lastModified "
                                     "queries. Only pages changed since last index are re-embedded.",
                                     "After team members update wiki pages — keeps answers current."),
    ("Clear Chat",        "Grey",    "Clears all messages from the chat window on-screen only. "
                                     "Does NOT delete the vector database or conversation memory.",
                                     "Start a fresh conversation topic without re-indexing."),
]

tbl_btn = doc.add_table(rows=len(btn_rows), cols=4)
tbl_btn.style = "Table Grid"
tbl_btn.alignment = WD_TABLE_ALIGNMENT.LEFT
col_w_btn = [Inches(1.15), Inches(0.6), Inches(2.9), Inches(1.85)]
for r_i, row_data in enumerate(btn_rows):
    row = tbl_btn.rows[r_i]
    for c_i, (txt, w) in enumerate(zip(row_data, col_w_btn)):
        cell = row.cells[c_i]
        cell.width = w
        if r_i == 0:
            set_cell_bg(cell, "3B4FD8")
        elif c_i == 1 and r_i > 0:
            val = row_data[1].lower()
            bg = "DDE8FF" if val == "blue" else "F0F0F0"
            set_cell_bg(cell, bg)
        elif c_i == 3 and r_i > 0:
            set_cell_bg(cell, "FFFDE8")
        elif r_i % 2 == 0:
            set_cell_bg(cell, "F5F7FF")
        else:
            set_cell_bg(cell, "FAFAFA")

        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        p.paragraph_format.left_indent  = Pt(5)
        p.paragraph_format.right_indent = Pt(3)
        run = p.add_run(txt)
        run.font.size = Pt(9)
        if r_i == 0:
            run.bold = True
            run.font.color.rgb = WHITE
        elif c_i == 0:
            run.bold = True
            run.font.color.rgb = BRAND_BLUE

doc.add_paragraph()

# ── 5.4 Progress Bar ──────────────────────────────────────────────────────────
heading(doc, "5.4  Progress Bar", 2)

callout(doc, [
    "Indexing Jira...   ████████████████░░░░░   78%   (55,780 / 71,385 chunks embedded)",
], bg="1A1A2E", fg=RGBColor(0xCC, 0xCC, 0xCC))

para(doc,
    "While any Index or Update button is running, a progress bar appears below the toolbar. "
    "It polls /api/ingest/progress every second and shows chunks_done / chunks_total. "
    "The bar disappears automatically when indexing completes and the status badge updates.",
    size=10.5, after=10)

# ── 5.5 Chat Window ────────────────────────────────────────────────────────────
heading(doc, "5.5  Chat Window", 2)

bullet(doc, "Answers stream in real-time token-by-token as GPT-4.1 generates them — you see the text forming live.")
bullet(doc, "Markdown is fully rendered: bold text, bullet lists, numbered steps, inline code, code blocks.")
bullet(doc, "Every answer shows a Sources section listing the exact documents, Jira IDs, and Confluence pages used.")
bullet(doc, "Confluence page citations appear as clickable hyperlinks — click to open the source page directly in your wiki.")
bullet(doc, "A Stop button appears during streaming — click it to cancel the response mid-generation.")
bullet(doc, "The AI remembers the conversation within a session — follow-up questions work correctly.")

doc.add_paragraph()

# ── 5.6 Input Box ──────────────────────────────────────────────────────────────
heading(doc, "5.6  Input Box", 2)

bullet(doc, "Type your question in plain English — no special syntax needed.")
bullet(doc, "Press Enter to send the question.")
bullet(doc, "Press Shift + Enter to add a new line without sending.")
bullet(doc, "The input box auto-expands as you type longer questions.")
bullet(doc, "The Send button is disabled while a response is streaming to prevent duplicate submissions.")

doc.add_paragraph()
doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 6 — SAMPLE ANSWERS
# ══════════════════════════════════════════════════════════════════════════════

heading(doc, "6.  Sample Questions & What ArcMind Returns", 1)

para(doc,
    "These examples show the types of questions your team can ask and the quality "
    "of answers ArcMind provides — sourced from your actual documentation, Jira, and Confluence.",
    size=11, after=10)

qa_examples = [
    (
        "Documentation question",
        "How do I configure SFTP authentication using a private key in Arc?",
        [
            "To configure key-based SFTP authentication:",
            "  1. Set AuthScheme to PublicKey in the SFTP Connector settings.",
            "  2. Set SSHClientCert to the file path of your .pem or .ppk key.",
            "  3. Set SSHClientCertType to PEMKEY_FILE (or PPKFILE for PuTTY format).",
            "  4. Set SSHClientCertPassword if your key has a passphrase.",
            "  5. Optionally set SSHServerFingerprint to prevent MITM attacks.",
            "",
            "  Sources: SFTP.html  ·  SFTP-Advanced-Properties.html",
        ],
        "Saves 20+ minutes of manual doc search. Returns exact property names with steps."
    ),
    (
        "Jira bug history question",
        "Has there been a report of the REST connector timing out on large payloads?",
        [
            "Yes — ARC-4821 (Jan 2025, Resolved): timeout on requests over 10 MB.",
            "Root cause: default Timeout property was 30s. Fix: increase Timeout to 120",
            "and enable ChunkSize streaming.",
            "",
            "ARC-5103 (Feb 2025, Open): related follow-up requesting MaxResponseSize property.",
            "",
            "Sources: ARC-4821  ·  ARC-5103",
        ],
        "Team immediately knows this was already reported and how it was resolved — no Jira search needed."
    ),
    (
        "Cross-source hybrid question",
        "Are there any known issues with FTP passive mode? How do I configure it?",
        [
            "Configuration (from FTP.html):",
            "  Set ConnectionType to PASSIVE. Recommended when behind a firewall.",
            "",
            "Known Issues (from Jira):",
            "  ARC-3317 — PASV response returns external IP. Workaround: UseRemoteAddress=false",
            "  ARC-3891 — Passive mode handshake timeout on slow networks. Fix in v22.3:",
            "             set ConnectTimeout=60",
            "",
            "Sources: FTP.html  ·  ARC-3317  ·  ARC-3891",
        ],
        "One question returns both the how-to config AND all known bugs — from two different sources."
    ),
]

for i, (category, question, answer_lines, why) in enumerate(qa_examples, 1):
    heading(doc, f"6.{i}  {category}", 2)

    callout(doc, ["Question", "", "  " + question], bg="E8ECFD")
    callout(doc, ["ArcMind Answer", ""] + ["  " + l for l in answer_lines], bg="E8F7EF")

    wp = doc.add_paragraph()
    wp.paragraph_format.space_after = Pt(12)
    wr = wp.add_run("Why this helps:  ")
    wr.bold = True
    wr.font.color.rgb = GREEN_ACCENT
    wr.font.size = Pt(10.5)
    wr2 = wp.add_run(why)
    wr2.font.size   = Pt(10.5)
    wr2.font.italic = True

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 7 — GITHUB REPOSITORY
# ══════════════════════════════════════════════════════════════════════════════

heading(doc, "7.  GitHub Repository", 1)

para(doc,
    "The full source code, README, Dockerfile, and setup guide are on GitHub:",
    size=11, after=12)

tbl_gh = doc.add_table(rows=1, cols=1)
tbl_gh.alignment = WD_TABLE_ALIGNMENT.CENTER
gh_cell = tbl_gh.cell(0, 0)
set_cell_bg(gh_cell, "1A1A2E")
gh_cell.width = Inches(5.5)
ghp = gh_cell.paragraphs[0]
ghp.alignment = WD_ALIGN_PARAGRAPH.CENTER
ghp.paragraph_format.space_before = Pt(18)
ghp.paragraph_format.space_after  = Pt(18)
ghp.paragraph_format.left_indent  = Pt(16)
ghp.paragraph_format.right_indent = Pt(16)
ghr1 = ghp.add_run("https://github.com/roopeshg123/arcmind")
ghr1.font.size      = Pt(14)
ghr1.font.bold      = True
ghr1.font.color.rgb = RGBColor(0x7C, 0x8F, 0xF5)
ghr1.font.name      = "Courier New"

doc.add_paragraph()

para(doc, "Repository contents:", size=11, after=5)
bullet(doc, "Full Python source code  — FastAPI backend + full RAG pipeline")
bullet(doc, "Dockerfile + docker-compose.yml  — one-command deployment")
bullet(doc, "README.md  — complete setup guide (Docker and manual)")
bullet(doc, "requirements.txt  — all pinned dependencies")
bullet(doc, "static/index.html  — complete chat UI, no build step needed")
bullet(doc, "PROJECT_ARCHITECTURE.txt  — deep-dive on every component")
bullet(doc, ".env template  — all configurable environment variables documented")

doc.add_paragraph()
divider(doc)
doc.add_paragraph()

# Final line
fp = doc.add_paragraph()
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = fp.add_run("ArcMind v2.1.0  ·  March 2026  ·  Built with OpenAI GPT-4.1 + RAG")
fr.font.size      = Pt(9)
fr.font.color.rgb = MID_GREY

# ─── Save ────────────────────────────────────────────────────────────────────
out_path = r"D:\arcmind\ArcMind_Project_Overview.docx"
doc.save(out_path)
print(f"Saved -> {out_path}")

